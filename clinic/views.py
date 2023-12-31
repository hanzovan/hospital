import json
from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.db import IntegrityError
from django.db.models import Q
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, timedelta, date
from django.utils import timezone

# for word document editing
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# for thousand separator
import locale

from .models import User, People, ContactDiary, Company, Service, Contract, MeetUp, MeetingAgendaItem
from .helpers import strong_password, user_right, days_between


# Create your views here.
@login_required
def index(request):
    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')

    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    # Get all the service
    services = Service.objects.all().order_by('name')

    return render(request, "clinic/index.html", {
        "yay_message": yay_message,
        "nay_message": nay_message,
        "services": services
    })


# Allow user to create their own account
def register(request):
    # If user submitted form
    if request.method == 'POST':
        # Define variables
        username = request.POST['username']
        email = request.POST['email']
        password = request.POST['password']
        confirm = request.POST['confirm']

        # Handle user input
        if not username or not email or not password or not confirm:
            return render(request, 'clinic/register.html', {
                'nay_message': 'Please fill all fields'
            })

        if password != confirm:
            return render(request, "clinic/register.html", {
                "nay_message": "Passwords do not match"
            })
        
        if not strong_password(password):
            return render(request, "clinic/register.html", {
                "nay_message": 'Password is not strong enough'
            })

        # If user input is okay, try to create a new account
        try:
            user = User.objects.create_user(username=username, email=email, password=password, management_right_level=0)
            login(request, user)
            request.session['yay_message'] = 'Registered successfully'
            return HttpResponseRedirect(reverse('index'))

        except IntegrityError:
            return render(request, "clinic/register.html", {
                "nay_message": "Username already taken"
            })

    # If user clicking link
    else:
        # Render the register page with register form
        return render(request, "clinic/register.html")


# Allow user to logged in with their account
def login_view(request):
    # User submitting form
    if request.method == 'POST':
        # Handle user input
        username = request.POST['username']
        password = request.POST['password']

        if not username or not password:
            return render(request, "clinic/login.html", {
                "nay_message": "Please fill all fields"
            })
        
        # authenticate user
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            request.session['yay_message'] = 'Logged in successfully'
            return HttpResponseRedirect(reverse('index'))

        else:            
            return render(request, "clinic/login.html", {
                "nay_message":"Invalid credentials"
            })

    # If user clicking link or being redirected
    else:
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')

        request.session['yay_message'] = ''
        request.session['nay_message'] = ''
        return render(request, "clinic/login.html", {
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Allow user to log out
@login_required
def logout_view(request):
    logout(request)
    request.session['yay_message'] = 'Logged out'
    return HttpResponseRedirect(reverse('index'))


# Allow client side to check the user right
@login_required
@csrf_exempt
def check_right(request):
    if request.method != 'POST':
        return JsonResponse({'error':'POST method required'}, status=400)
    data = json.loads(request.body)
    if data.get('right') is not None:
        right = data['right']
        check_result = (right in user_right(request.user.management_right_level))
        return JsonResponse({'check_result': check_result}, safe=False)

    else:
        return JsonResponse({'message': 'no info'})

# Allow top level admin to manage to list of users, and to apply management level for other users
@login_required
@csrf_exempt
def authorize(request):
    # if admin submitted form
    if request.method == 'POST':
        # User must have the right to submit form
        if 'modify_user_right' not in user_right(request.user.management_right_level):
            return JsonResponse({'error': "You do not have rights to change user's level"}, status=400)

        data = json.loads(request.body)

        if data.get('new_level') is not None and data.get('chosen_user_id') is not None:
            # Get the user
            user = User.objects.get(pk=data['chosen_user_id'])
            user.management_right_level = data['new_level']
            user.save()            
            
            return JsonResponse({'message': 'Level changed'})
        else:
            return JsonResponse({'message': 'Nothing changed'})
    
    # if admin clicking link
    else:
        # if user does not have right to access this page
        if 'read_user_right' not in user_right(request.user.management_right_level):            
            request.session['nay_message'] = "You don't have the right to visit this page"
            return HttpResponseRedirect(reverse('index'))
        modify_right = 'modify_user_right' in user_right(request.user.management_right_level)
        data_users = User.objects.all()
        return render(request, "clinic/authorize.html", {
            "data_users": data_users,
            "levels": [1, 2, 3],
            "modify_right": modify_right
        })
    

# Allow user to add service
# Check if user have right to add service or not
@login_required
def add_service(request):
    # If user submitted form
    if request.method == 'POST':
        # If user do not have right, raise error
        if 'modify_service_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = 'You do not have right to do this'
            return HttpResponseRedirect(reverse('index'))

        # Define variables
        name = request.POST['name']
        male_price = request.POST.get('male_price')
        female_price = request.POST.get('female_price')
        benefit = request.POST.get('benefit', '')
        description = request.POST.get('description', '')

        if not name or (not male_price and not female_price):
            return render(request, "clinic/add_service.html", {
                "nay_message": "Please fill at least name, price of the service for male or female"
            })

        if not male_price:
            male_price = None
        if not female_price:
            female_price = None
        
        old_services = Service.objects.filter(Q(name__icontains=name))            
        if old_services.count() > 0:
            return render(request, "clinic/add_service.html", {
                "nay_message": f"{name} already in database"
            })

        else:
            # if user input correct, add to the services
            service = Service(
                name = name,
                male_price = male_price,
                female_price = female_price,
                benefit = benefit,
                description = description,
                created_by = request.user
            )
            service.save()
            request.session['yay_message'] = 'Service added'
            return HttpResponseRedirect(reverse('index'))

    # If user clicked link
    else:
        # Only user with the modify service right can visit this page
        if 'modify_service_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = 'You do not have right to access this part'
            return HttpResponseRedirect(reverse('index'))
        
        return render(request, "clinic/add_service.html")


# Allow user to check the service's detail
@login_required
@csrf_exempt
def service_detail(request, service_id):
    # If user is admin and submited the form
    if request.method == 'POST':
        # If user does not have the right to edit, response error by json, then use JS to redirect user to index page
        if 'modify_service_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = 'You do not have the right to modify this information'
            return JsonResponse({'error':'You do not have the right to edit'}, status=403)
        data = json.loads(request.body)

        if data.get('new-name') is not None:
            try:
                service = Service.objects.get(pk=service_id)
                service.name = data['new-name']

                if not data.get('new-male-price'):
                    service.male_price = None
                else:
                    service.male_price = data['new-male-price']

                if not data.get('new-female-price'):
                    service.female_price = None
                else:
                    service.female_price = data['new-female-price']
                
                service.benefit = data['new-benefit']
                service.description = data['new-description']
                service.modified_by = request.user
                service.save()
                return JsonResponse({'message': 'Changes saved'}, status=200)
            except Service.DoesNotExist:
                return JsonResponse({'error': 'Service does not exist in database'}, status=400)

    else:
        if 'read_all_service_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to check service's detail"
            return HttpResponseRedirect(reverse('index'))
        try:
            service = Service.objects.get(pk=service_id)

            yay_message = request.session.get('yay_message', '')
            nay_message = request.session.get('nay_message', '')
            request.session['yay_message'] = ''
            request.session['nay_message'] = ''

            return render(request, "clinic/service_detail.html", {
                "service": service,
                "yay_message": yay_message,
                "nay_message": nay_message
            })
        except Service.DoesNotExist:
            request.session['nay_message'] = 'Service does not exist'
            return HttpResponseRedirect(reverse('index'))


# Allow user to remove service
@login_required
def remove_service(request):
    if request.method == 'POST':
        # Check user right
        if 'modify_service_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to modify service information"
            return HttpResponseRedirect(reverse('index'))

        # Get the service that needed to remove
        service_id = request.POST.get('service_id', '')

        try:
            service = Service.objects.get(pk=service_id)
        except Service.DoesNotExist:
            request.session['nay_message'] = "Service with provided id does not exist"
        
        service.delete()

        request.session['yay_message'] = "Service removed"

        return HttpResponseRedirect(reverse('index'))
    
    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('index'))


# Allow user to add people
@login_required
def add_people(request):
    # If user submitted form
    if request.method == 'POST':
        # Check if user have the right, it no, redirect to index, if yes then continue
        if 'add_people_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add person's information"
            return HttpResponseRedirect(reverse('index'))

        # Handle user input
        name = request.POST['name']
        position = request.POST.get('position', '')
        address = request.POST.get('address', '')
        email = request.POST.get('email', '')
        phone = request.POST.get('phone', '')
        note = request.POST.get('note', '')

        if not name:
            return render(request, "clinic/add_people.html", {
                "nay_message": "Name required"
            })

        # Check if person has duplicate information
        name_duplicate = People.objects.filter(name=name).exists()
        email_duplicate = People.objects.filter(email=email).exists()
        phone_duplicate = People.objects.filter(phone=phone).exists()
        if name_duplicate or email_duplicate or phone_duplicate:
            request.session['nay_message'] = "Duplicate information: check name, email, phone"
            return HttpResponseRedirect(reverse('add_people'))

        # If user input is okay, trying to create new people in database        
        people = People(
            name = name,
            position = position,
            address = address,
            email = email,
            phone = phone,
            note = note,
            created_by = request.user
        )
        people.save()

        request.session['yay_message'] = 'People saved'

        # If user has the permission to read people information, redirect to people page, else to the people added by me page
        if "read_all_people_info" not in user_right(request.user.management_right_level):
            return HttpResponseRedirect(reverse('my_people'))
        else:
            return HttpResponseRedirect(reverse('people'))        

    # If user clicked link:
    else:
        if 'add_people_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = 'You do not have the right to add person information'
            return HttpResponseRedirect(reverse('index'))

        # Get the companies
        companies = Company.objects.all()

        # Get message
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        return render(request, "clinic/add_people.html",{
            "companies": companies,
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Allow user to check people added by current user
@login_required
def my_people(request):
    if "read_self_add_people_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You're not allow to check this info"
        return HttpResponseRedirect(reverse('index'))
    
    # Get message if there is
    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    # Get people information
    people = People.objects.filter(created_by=request.user)

    # Get latest message as well as the companies that related to the person
    for person in people:
        person.latest_message = person.talks.order_by("-date").first()

        # Get the list of companies
        # related_companies = person.company.all()
        # representing_companies = person.companies.all()

        # unique_companies = list(set(related_companies).union(set(representing_companies)))
        # person.unique_companies = unique_companies

    return render(request, "clinic/my_people.html", {
        "people": people,
        "yay_message": yay_message,
        "nay_message": nay_message
    })


# Allow user to check all people relevant to the hospital, only level 2 or higher admin can check this
@login_required
def people(request):
    # Check if user have appropriate right to visit the page
    if "read_all_people_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = 'You are not allowed to enter this part'
        return HttpResponseRedirect(reverse('index'))

    nay_message = request.session.get('nay_message', '')
    yay_message = request.session.get('yay_message', '')

    request.session['nay_message'] = ''    
    request.session['yay_message'] = ''

    people = People.objects.all()

    for person in people:
        # Get latest message of the person
        person.latest_message = person.talks.order_by("-date").first()

        # These lines are use to create a temporary field (not the normal field in models.py) for person instance, that including all unique company that related to the person, it was created by compare 2 sets, then choose the unique instance
        # People used to have a field named company that related with people by ManyToMany relationship

        # related_companies = person.company.all()

        # Get the list of company that choose this person as the representative
        # representing_companies = person.companies.all()

        # unique_companies = list(set(related_companies).union(set(representing_companies)))
        # person.unique_companies = unique_companies

    return render(request, "clinic/people.html", {
        "people": people,
        "nay_message": nay_message,
        "yay_message": yay_message
    })


# Allow admin to check people's information and modify it
@login_required
@csrf_exempt
def person_detail(request, person_id):
    # If user submited the form
    if request.method == 'POST':
        # Check if user have the right to modify person's info
        if "modify_people_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to modify this information"
            return JsonResponse({'error': 'You do not have the right to modify this info'}, status=403)
        
        # Get variables
        data = json.loads(request.body)

        # If user do not leave empty name for person, update new information
        if data.get('new_name') is not None:
            try:
                person = People.objects.get(pk=person_id)
                person.name = data['new_name']
                person.position = data['new_position']
                person.address = data['new_address']
                person.email = data['new_email']
                person.phone = data['new_phone']
                person.note = data['new_note']
                person.modified_by = request.user
                
                person.save()

                return JsonResponse({'message': 'New info saved'}, status=200)

            except People.DoesNotExist:
                request.session['nay_message'] = 'That person does not exist in database'
                return JsonResponse({'error': 'Person does not exist'}, status=400)

        # If user leave empty name, return error
        else:
            request.session['nay_message'] = "Don't leave empty name"
            return JsonResponse({'error': "Don't leave empty name"}, status=401)

    # If admin clicked link or visit the correct url
    else:
        # If user doesn't even have the right to read self add people, they should be redirected to index page
        if "read_self_add_people_info" not in user_right(request.user.management_right_level):
            return HttpResponseRedirect(reverse('index'))        

        # If user is valid for the right, continue
        try:
            person = People.objects.get(pk=person_id)            
        except People.DoesNotExist:
            request.session['nay_message'] = 'That person does not exist'

            # Redirect base on user permission if person not found
            if "read_all_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('people'))
            else:
                return HttpResponseRedirect(reverse('my_people'))
        
        # If person found, check if that person was added by that user, if not user need to have permission to read all people info or will being redirected to index page
        if "read_all_people_info" not in user_right(request.user.management_right_level) and person.created_by != request.user:
            request.session['nay_message'] = 'You do not have the right to visit this page'
            return HttpResponseRedirect(reverse('index'))
        
        # Get the message from this person
        messages = person.talks.all().order_by("-date")

        # Get messages from session
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        return render(request, "clinic/person_detail.html", {
            "person": person,
            "messages": messages,
            'yay_message': yay_message,
            'nay_message': nay_message
        })


# Allow user with appropriate right to remove person's information
@login_required
def remove_person(request):
    if request.method == 'POST':
        # Check if user have the appropriate right
        if 'modify_people_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to do this"
            return HttpResponseRedirect(reverse('index'))

        # Define variable
        person_id = request.POST.get('person_id', '')
        try:
            person = People.objects.get(pk=person_id)
        except People.DoesNotExist:
            request.session['nay_message'] = "People with that id does not exist"
            return HttpResponseRedirect(reverse('people'))
        
        # If user is legit, continue
        person.delete()
        request.session['yay_message'] = "Person information deleted successfully"
        return HttpResponseRedirect(reverse('people'))
    
    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('index'))


# Add company
@login_required
def add_company(request):
    # If user submit form
    if request.method == 'POST':
        # Check if user have right
        if "add_company_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You don't have the right to add company"
            return HttpResponseRedirect(reverse('index'))

        # If user have right, continue to check user input
        name = request.POST['name']
        industry = request.POST['industry']
        address = request.POST['address']
        email = request.POST.get('email', '')
        phone = request.POST.get('phone', '')
        male_headcount = request.POST['male_headcount']
        female_headcount = request.POST['female_headcount']

        # Required user to fill information
        if not name or not industry or not address or not male_headcount or not female_headcount or not email or not phone:
            return render(request, "clinic/add_company.html", {
                "nay_message": "All fields required"
            })
        
        # Check duplicate name, email, phone
        duplicate_name = Company.objects.filter(name=name).exists()
        duplicate_email = Company.objects.filter(email=email).exists()
        duplicate_phone = Company.objects.filter(phone=phone).exists()

        if duplicate_name or duplicate_email or duplicate_phone:
            request.session['nay_message'] = "Duplicate informations, check name, email, or phone"
            return HttpResponseRedirect(reverse('add_company'))
        
        try:
            if int(male_headcount) < 0 or int(female_headcount) < 0:
                nay_message = "Please enter positive number"
                return render(request, "clinic/add_company.html", {
                    "nay_message": nay_message
                })
        except ValueError:
            request.session['nay_message'] = "Headcount has to be integer"
            return redirect("add_company")
        
        # Save the company        
        new_company = Company(
            name = name,
            industry = industry,
            address = address,
            email = email,
            phone = phone,
            male_headcount = male_headcount,
            female_headcount = female_headcount,
            created_by = request.user
        )

        # Get the method that user choose to add company's representative
        representative_add_method = request.POST.get('add_representative_method', '')

        # If user decide to add representative method
        if representative_add_method:
            # If user decide to add representative as a new person
            if representative_add_method == 'create_new':
                representative_name = request.POST.get('new_representative_name', '')
                representative_address = request.POST.get('new_representative_address', '')
                representative_email = request.POST.get('new_representative_email', '')
                representative_phone = request.POST.get('new_representative_phone', '')

                if not representative_name:
                    request.session['nay_message'] = "New representative name missing"
                    return HttpResponseRedirect(reverse('add_company'))

                # Check duplication
                name_duplicate = People.objects.filter(name=representative_name).exists()
                email_duplicate = People.objects.filter(email=representative_email).exists()
                phone_duplicate = People.objects.filter(phone=representative_phone).exists()
                if name_duplicate or email_duplicate or phone_duplicate:
                    request.session['nay_message'] = "Duplicate information: Check representative's name, email, phone"
                    return HttpResponseRedirect(reverse('add_company'))

                # create new person information
                representative = People(
                    name = representative_name,
                    address = representative_address,
                    email = representative_email,
                    phone = representative_phone,
                    created_by = request.user
                )
                representative.save()
                new_company.representative = representative

            # If user decide to add representative from people list
            elif representative_add_method == 'choose_from_list':
                representative_id = request.POST.get('representative_id', '')

                if not representative_id:
                    request.session['nay_message'] = "Representative was not chosen"
                    return HttpResponseRedirect(reverse('add_company'))
                try:
                    representative = People.objects.get(pk=representative_id)
                except People.DoesNotExist:
                    request.session['nay_message'] = "Person not found"
                    return HttpResponseRedirect(reverse('add_company'))
                
                # Save the representative to the company
                new_company.representative = representative

            # else if other method, ignore it

        # Save the company information
        new_company.save()

        request.session['yay_message'] = "Company added"
        return HttpResponseRedirect(reverse('companies'))            
    
    # If user clicking link or being redirected
    else:
        if "add_company_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to visit this part"
            return HttpResponseRedirect(reverse('index'))
        
        # Get people list
        people = People.objects.all()

        yay_message = request.session.get('yay_messgae', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['yay_message'] = ''

        return render(request, "clinic/add_company.html", {
            "people": people,
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Allow user to add company's information
@login_required
def companies(request):
    # Only user with the right to check company info can go to this route
    if "read_company_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the right to access this page"
        return HttpResponseRedirect(reverse('index'))
    # If user right satisfied condition, show all companies
    companies = Company.objects.all()

    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    return render(request, "clinic/companies.html", {
        "companies": companies,
        "yay_message": yay_message,
        "nay_message": nay_message
    })


# Allow user to access company's detail
@login_required
def company_detail(request, company_id):
    # If user submitted form
    if request.method == 'POST':
        # Check the id
        try:
            company = Company.objects.get(pk=company_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Company not found"

            # If user does not have permission to read all companies information, redirect to index page, else redirect to company page
            if "read_company_info" not in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('index'))
            else:
                return HttpResponseRedirect(reverse('companies'))

        # If the id is valid, then company exist, check the user right
        if "modify_company_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to modify company's information"

            # If user has permission to check person detail, redirect to that page, else to the index page
            if "read_company_info" not in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('index'))
            else:
                return redirect("company_detail", company_id=company_id)
        
        # Get the variables
        name = request.POST.get('name', '')
        industry = request.POST.get('industry', '')
        address = request.POST.get('address', '')
        email = request.POST.get('email', '')
        phone = request.POST.get('phone', '')
        male_headcount = request.POST.get('male_headcount', '')
        female_headcount = request.POST.get('female_headcount', '')

        if not name or not industry or not address or not email or not phone or not male_headcount or not female_headcount:
            request.session['nay_message'] = "Please fill all fields"
            return redirect("company_detail", company_id=company_id)

        try:
            # Check that if headcount is negative or not
            if int(male_headcount) < 0 or int(female_headcount) < 0:
                request.session['nay_message'] = "Headcount has to be equal or greater than 0"
                return redirect("company_detail", company_id=company_id)
        except ValueError:
            request.session['nay_message'] = "Headcount has to be integer"
            return redirect("company_detail", company_id=company_id)        

        # Save the new information of the company
        company.name = name
        company.industry = industry
        company.address = address
        company.email = email
        company.phone = phone
        company.male_headcount = male_headcount
        company.female_headcount = female_headcount
        company.modified_by = request.user
        company.save()

        # Get the variables for update representatives
        update_method = request.POST.get('update_representative_method', '')

        # Decide the approach based on how user chose the method, if not choose one of the 2, representative will not be updated
        if update_method:
            if update_method == "create_new":
                new_representative_name = request.POST.get('new_representative_name', '')
                new_representative_address = request.POST.get('new_representative_address', '')
                new_representative_email = request.POST.get('new_representative_email', '')
                new_representative_phone = request.POST.get('new_representative_phone', '')

                # Check if there is duplicate in name, email, or phone
                name_duplicate = People.objects.filter(name=new_representative_name).exists()
                email_duplicate = People.objects.filter(email=new_representative_email).exists()
                phone_duplicate = People.objects.filter(phone=new_representative_phone).exists()
                if name_duplicate or email_duplicate or phone_duplicate:
                    request.session['nay_message'] = "Update representative failed: Duplicate information, check name, email, and phone"
                    return redirect('company_detail', company_id=company_id)

                representative = People(
                    name = new_representative_name,
                    address = new_representative_address,
                    email = new_representative_email,
                    phone = new_representative_phone,
                    created_by = request.user
                )
                representative.save()
                company.representative = representative
                company.save()
            
            elif update_method == "choose_from_list":
                representative_id = request.POST.get('representative_id', '')

                # If user somehow did not submit the id, return error
                if not representative_id:
                    request.session['nay_message'] = "Update representative failed: Person id not selected"
                    return redirect('company_detail', company_id=company_id)
                
                # Check if id is valid
                try:
                    representative = People.objects.get(pk=representative_id)
                except People.DoesNotExist:
                    request.session['nay_message'] = "Update representative failed: Person not found"
                    return redirect('company_detail', company_id=company_id)

                company.representative = representative
                company.save()

            #else: ignore the representative part

        # Inform user and redirect
        request.session['yay_message'] = "Company information modified successfully"
        return redirect('company_detail', company_id=company_id)
    
    # If user click link or being redirected
    else:
        # Only user with the right to check company info can go to this route
        if "read_company_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to access this page"
            return HttpResponseRedirect(reverse('index'))
        # If user right satisfied condition, continue
        try:
            company = Company.objects.get(pk=company_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Company with that id does not exist"
            return HttpResponseRedirect(reverse('companies'))

        # Get people list
        people = People.objects.all()

        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        return render(request, "clinic/company_detail.html", {
            "company": company,
            "people": people,
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Allow user to remove company
@login_required
def remove_company(request):
    if request.method == 'POST':
        company_id = request.POST.get('company_id', '')
        try:
            company = Company.objects.get(pk=company_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Company with that id does not exist"

            # If user has permission to read companies info
            if "read_company_info" in user_right(request.user.management_right_level):                
                return HttpResponseRedirect(reverse('companies'))
            else:
                return HttpResponseRedirect(reverse('index'))

        # Check user right
        if 'modify_company_info' not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to modify company information"

            # If user has permission to read company info
            if "read_company_info" in user_right(request.user.management_right_level):
                return redirect('company_detail', company_id=company_id)
            else:
                return HttpResponseRedirect(reverse('index'))

        # Remove company
        company.delete()

        # Alert success and redirect
        request.session['yay_message'] = "Company was deleted"
        return HttpResponseRedirect(reverse('companies'))
    
    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('index'))


# Allow add message from person_detail page
def message(request, person_id):
    if request.method == 'POST':
        #Check if person with that id exist or not, if not redirect user to index page
        try:
            person = People.objects.get(pk=person_id)
        except People.DoesNotExist:
            request.session['nay_message'] = "Person not found"

            # Check user permission and redirect depend on that
            if "read_all_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('people'))
            elif "read_self_add_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('my_people'))
            else:
                return HttpResponseRedirect(reverse('index'))

        #Check user right
        if "add_people_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add message"

            # Check user permission and redirect
            if "read_all_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('people'))
            elif "read_self_add_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('my_people'))
            else:
                return HttpResponseRedirect(reverse('index'))

        #Get content from the form input
        content = request.POST.get('content', '')

        if not content:
            request.session['nay_message'] = "Message's content was empty, no message saved"
            return redirect('person_detail', person_id=person.id)

        message = ContactDiary(
            name = person,
            content = content
        )
        message.save()
        request.session['yay_message'] = "Message saved"
        return redirect('person_detail', person_id=person.id)

    else:
        request.session['nay_message'] = "POST request required"
        return HttpResponseRedirect(reverse('index'))

# Add message from all people page, with symbol for better visual
@login_required
def add_message_from_all_people_page(request, person_id):
    # If user submitted form
    if request.method == 'POST':
        #Check if person exist
        try:
            person = People.objects.get(pk=person_id)
        except People.DoesNotExist:
            request.session['nay_message'] = "Person does not exist"

            # Redirect user base on user permission
            if "read_all_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('people'))
            elif "read_self_add_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('my_people'))
            else:
                return HttpResponseRedirect(reverse('index'))
        
        #Check if user has the right, if not raise error and show the form again
        if "add_people_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add message"
            return HttpResponseRedirect(reverse('index'))

        content = request.POST.get('content', '')
        if not content:
            request.session['nay_message'] = "Please fill the message's content"
            return redirect("add_message_from_all_people_page", person_id=person_id)

        # save the new message
        new_message = ContactDiary(
            name = person,
            content = content
        )
        new_message.save()
        request.session['yay_message'] = "Message saved"

        # Redirect base on user permission
        if "read_all_people_info" in user_right(request.user.management_right_level):
            return HttpResponseRedirect(reverse('people'))
        else:
            return HttpResponseRedirect(reverse('my_people'))
    
    # If user clicked link
    else:
        # Check if user has the right, if not redirect user to index route
        if "add_people_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add message"
            return HttpResponseRedirect(reverse('index'))

        # Check if person exist
        try:            
            person = People.objects.get(pk=person_id)
        except People.DoesNotExist:
            request.session['nay_message'] = "Person not found"

            # Redirect base on user permission
            if "read_all_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('people'))
            elif "read_self_add_people_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('my_people'))
            else:
                return HttpResponseRedirect(reverse('index'))            
        
        # Show message
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        return render(request, "clinic/add_message_from_all_people_page.html", {
            "person": person,
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Add contract
@login_required
def add_contract(request):
    # If user submitting form
    if request.method == 'POST':
        # Check if user has the right to modify contract info
        if "modify_contract_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add contract"
            return HttpResponseRedirect(reverse('index'))

        # Get the client id
        client_id = request.POST.get('client_id', '')

        if not client_id:
            request.session['nay_message'] = "Choose a company"
            return HttpResponseRedirect(reverse('add_contract'))

        # Try to get the company, if there ain't that company, raise error
        try:
            client = Company.objects.get(pk=client_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Company does not exist"
            return HttpResponseRedirect(reverse('add_contract'))

        # Company has to provide representative before creating contract
        if not client.representative:
            request.session['nay_message'] = "Company has to update representative before creating contract"
            return HttpResponseRedirect(reverse('add_contract'))
        
        # Get service id
        service_ids = request.POST.getlist('chosen_services')

        # Try to get the service, if there ain't that service, ignore that
        services = []
        for i in service_ids:
            try:
                service = Service.objects.get(pk=i)            
                services.append(service)
            except Service.DoesNotExist:
                request.session['nay_message'] = f"service with id {i} does not exist"

        # Get headcount, but checked if they're positive integer        
        male_headcount = request.POST.get('male_headcount', '')
        female_headcount = request.POST.get('female_headcount', '')

        try:
            male_number = int(male_headcount)
            female_number = int(female_headcount)
            if male_number < 0 or female_number < 0:
                request.session['nay_message'] = "headcount has to be positive integer, not negative"
                return HttpResponseRedirect(reverse('add_contract'))
        except ValueError:
            request.session['nay_message'] = 'headcount has to be positive integer, not a float'
            return HttpResponseRedirect(reverse('add_contract'))
        
        # Get the total value
        total_value = 0
        for service in services:
            if service.male_price:
                service_value_for_male = int(service.male_price)*male_number
            else:
                service_value_for_male = 0
            
            if service.female_price:
                service_value_for_female = int(service.female_price)*female_number
            else:
                service_value_for_female = 0
            total_value += (service_value_for_male + service_value_for_female)
        
        # Get the discount
        discount = request.POST.get('discount')
        if discount:
            discount = round(float(discount), 2)
            if discount < 0:
                request.session['nay_message'] = "Discount has to be equal or greater than 0"
                return HttpResponseRedirect(reverse('add_contract'))        

        # Get the total revenue
        revenue = round(float(total_value*(100 - discount)/100))

        # Get the initiation date
        initiation_date = request.POST.get('initiation_date', '')   
        new_contract = Contract(
            client = client,
            male_headcount = male_headcount,
            female_headcount = female_headcount,
            total_value = total_value,
            discount = discount,
            revenue = revenue,
            initiation_date = initiation_date,
            created_by = request.user
        )
        new_contract.save()
        new_contract.services.set(services)
        new_contract.save()

        # Get the contract file
        pdf_file = request.FILES.get('contract_file')
        if pdf_file:
            # Create a new Contract instance with the uploaded file
            new_contract.pdf_file = pdf_file
            new_contract.save()

        request.session['yay_message'] = "Contract added"

        return HttpResponseRedirect(reverse('all_contracts'))
    
    # If user clicked link or being redirect
    else:
        # Check user right
        if "modify_contract_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add contract"

            # Redirect base on the user permission
            if "read_contract_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('all_contracts'))
            else:
                return HttpResponseRedirect(reverse('index'))
        
        companies = Company.objects.all()
        services = Service.objects.all()
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''
        return render(request, "clinic/add_contract.html", {
            "services": services,
            "companies": companies,
            "yay_message": yay_message,
            "nay_message": nay_message
        })


# Retrieve all contracts that are not archived
@login_required
def all_contracts(request):
    # If user does not has the right to read contract info, redirect to index
    if "read_contract_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = 'You do not have the right to access this information'
        return HttpResponseRedirect(reverse('index'))

    contracts = Contract.objects.filter(archived=False).order_by('initiation_date')

    # Get the current date
    today = date.today()

    # Add a field to each contract to confirm if the initiation date is near within 10 days
    for contract in contracts:
        if days_between(str(today), str(contract.initiation_date)) <= 10:
            contract.in_within_10_days = True
        else:
            contract.in_within_10_days = False

    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    return render(request, "clinic/contracts.html", {
        "contracts": contracts,
        "today": today,
        "yay_message": yay_message,
        "nay_message": nay_message
    })


# Allow user to access contract detail, and archive the contract
@login_required
def contract_detail(request, contract_id):
    # Check user right, if not redirect to index
    if "read_contract_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the right to access this page"
        return HttpResponseRedirect(reverse('index'))

    # If user has permission, check contract id
    try:
        contract = Contract.objects.get(pk=contract_id)
    except Contract.DoesNotExist:
        request.session['nay_message'] = "Invalid contract ID"

        # Redirect to all contracts page        
        return HttpResponseRedirect(reverse('all_contracts'))

    # List of companies for editing
    companies = Company.objects.all()
    services = Service.objects.all()
    chosen_services = contract.services.all()

    # Add field to service to define which services were chosen in original contract
    for service in services:
        if service in chosen_services:
            service.chosen = True
        else:
            service.chosen = False

    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    return render(request, "clinic/contract_detail.html", {
        "contract": contract,
        "companies": companies,
        "services": services,
        "yay_message": yay_message,
        "nay_message": nay_message
    })


# Allow user to edit the contract, delete the contract
@login_required
def edit_contract(request):
    if request.method == 'POST':
        # Check if user has the permission to edit contract information
        if "modify_contract_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the permission to modify this information"

            # Redirect base on user permission
            if "read_contract_info" in user_right(request.user.management_right_level):
                return HttpResponseRedirect(reverse('all_contracts'))
            else:
                return HttpResponseRedirect(reverse('index'))

        # Get variables from POST
        contract_id = request.POST.get('contract_id', '')
        client_id = request.POST.get('client_id', '')
        service_ids = request.POST.getlist('chosen_services')
        discount = request.POST.get('discount', '')
        male_headcount = request.POST.get('male_headcount', '')
        female_headcount = request.POST.get('female_headcount', '')
        initiation_date = request.POST.get('initiation_date', '')
        pdf_file = request.FILES.get('contract_file')

        # Check contract id
        if not contract_id:
            request.session['nay_message'] = "Contract id was not provided"
            return HttpResponseRedirect(reverse('all_contracts'))
        try:
            contract = Contract.objects.get(pk=contract_id)
        except Contract.DoesNotExist:
            request.session['nay_message'] = "Contract does not exist"
            return HttpResponseRedirect(reverse('all_contracts'))

        # Check if contract was archived
        if contract.archived:
            request.session['nay_message'] = "Archived contracts can't be edited"
            return redirect("contract_detail", contract_id=contract_id)

        # Check client id
        if not client_id:
            request.session['nay_message'] = "Please choose a client"
            return redirect("contract_detail", contract_id=contract_id)
        try:
            client = Company.objects.get(pk=client_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Company does not exist"
            return redirect("contract_detail", contract_id=contract_id)

        # Remove empty id in service_ids
        service_ids = [i for i in service_ids if i]

        # Check chosen service
        services = []
        for i in service_ids:
            try:
                service = Service.objects.get(pk=i)
                services.append(service)
            except Service.DoesNotExist:
                request.session['nay_message'] = f"service with id {i} does not exist"

        # Check discount
        if discount:
            try:
                discount = round(float(discount), 2)
                if discount < 0:
                    request.session['nay_message'] = "Discount has to be greater or equal to 0"
                    return redirect("contract_detail", contract_id=contract_id)
            except ValueError:
                request.session['nay_message'] = "Invalid input for discount"
                return redirect("contract_detail", contract_id=contract_id)
        else:
            discount = 0

        # Check male and female headcount
        # if both male_headcount and female_headcount are empty, raise error
        if not male_headcount and not female_headcount:
            request.session['nay_message'] = "Contract can not have empty head count"
            return redirect("contract_detail", contract_id=contract_id)
        
        # if either of male_headcount or female_headcount exist, move forward
        if not male_headcount:
            male_headcount = 0
        else:
            try: 
                male_headcount = int(male_headcount)
                if male_headcount < 0:
                    request.session['nay_message'] = "headcount has to be equal or greater than 0"
                    return redirect("contract_detail", contract_id=contract_id)
            except ValueError:
                request.session['nay_message'] = "headcount has to be positive integer"
                return redirect("contract_detail", contract_id=contract_id)

        if not female_headcount:
            female_headcount = 0
        else:
            try:
                female_headcount = int(female_headcount)
                if female_headcount < 0:
                    request.session['nay_message'] = "headcount has to be equal or greater than 0"
                    return redirect("contract_detail", contract_id=contract_id)
            except ValueError:
                request.session['nay_message'] = "headcount has to be positive integer"
                return redirect("contract_detail", contract_id=contract_id)
            
        # Check initiation date
        if not initiation_date:
            request.session['nay_message'] = "Initiation date required"
            return redirect("contract_detail", contract_id=contract_id)
        else:
            try:
                # Check the value of initiation_date
                initiation_date_check = datetime.strptime(initiation_date, '%Y-%m-%d')
            except ValueError:
                request.session['nay_message'] = "Invalid date format"
                return redirect("contract_detail", contract_id=contract_id)


        # Get the total value
        total_value = 0
        for service in services:
            if service.male_price:
                service_value_for_male = int(service.male_price)*male_headcount
            else:
                service_value_for_male = 0

            if service.female_price:
                service_value_for_female = int(service.female_price)*female_headcount
            else:
                service_value_for_female = 0
            total_value += (service_value_for_male + service_value_for_female)

        # Get the revenue
        revenue = round(float(total_value*(100 - discount)/100))

        # Update the contract value
        contract.client = client
        contract.male_headcount = male_headcount
        contract.female_headcount = female_headcount
        contract.total_value = total_value
        contract.discount = discount
        contract.revenue = revenue
        contract.initiation_date = initiation_date
        contract.modified_by = request.user
        contract.services.set(services)
        contract.save()

        # Get contract file
        if pdf_file:
            contract.pdf_file = pdf_file
            contract.save()

        request.session['yay_message'] = "Contract modified"
        return redirect("contract_detail", contract_id=contract_id)
    
    else:
        request.session['nay_message'] = "Post method required"
        return HttpResponseRedirect(reverse('all_contracts'))


# Allow user to change contracts from active to finish or archived
@login_required
def archive_contract(request):
    if request.method == 'POST':
        # Check if user has the permission to archive contract
        if "modify_contract_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the permission to modify contract information"
            return HttpResponseRedirect(reverse('all_contracts'))
        contract_id = request.POST.get('contract_id', '')
        if contract_id:
            # Check if that id match with any contract
            try:
                contract = Contract.objects.get(pk=contract_id)                
            except Contract.DoesNotExist:
                request.session['nay_message'] = "Contract with that id does not exist"
                return HttpResponseRedirect(reverse('all_contracts'))

            # If requirements met, archive the contract
            contract.archived = True
            contract.archived_date = timezone.now()
            contract.archived_by = request.user
            contract.save()

            # Inform successfully
            request.session['yay_message'] = "Contract archived successfully"
            return redirect('contract_detail', contract_id=contract_id)

        else:
            return HttpResponseRedirect(reverse('all_contracts'))
    
    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('all_contracts'))


# Allow user to visit all archived contracts
@login_required
def all_archived_contracts(request):
    #If user do not have the right to read contract info, redirect
    if "read_contract_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the permission to access this page"
        return HttpResponseRedirect(reverse('index'))

    # Get all archived contracts
    contracts = Contract.objects.filter(archived=True).order_by("initiation_date")

    return render(request, "clinic/all_archived_contracts.html",{
        "contracts": contracts
    })


# Prepare a contract in word ready for printing
@login_required
def generate_contract_docx(request, contract_id):
    #Check if user has the permission to read contract info
    if "read_contract_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the right to access this information"
        return HttpResponseRedirect(reverse('index'))

    # Get the contract
    contract = Contract.objects.get(pk=contract_id)

    # Create a new word document
    doc = Document()

    # Add content to the document
    title = doc.add_heading('CONTRACT AGREEMENT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)

    # Add contract detail to the document
    head1 = doc.add_heading('By King Hospital', level=3)
    head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head1_run = head1.runs[0]
    head1_run.font.size = Pt(16)

    # Date of Agreement
    contract_date = doc.add_paragraph(f"This Agreement is made on: {date.today()}")
    contract_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contract_date_run = contract_date.runs[0]
    contract_date_run.font.size = Pt(12)

    # Between
    doc.add_heading('BETWEEN')

    # Create 2 styles for indentation
    style = doc.styles.add_style('IndentedStyle', 1)
    style.font.size = Pt(12)
    style.paragraph_format.left_indent = Inches(0.5)

    style1 = doc.styles.add_style('IndentedStyle1', 1)
    style1.font.size = Pt(12)
    style1.paragraph_format.left_indent = Inches(0.2)

    # Add the indented paragraph with the defined style
    doc.add_heading('Party A: King Hospital', level=2)
    doc.add_paragraph("Address: 123 King street, ward 1, district 2, Ho Chi Minh city, Vietnam", style="IndentedStyle")
    doc.add_paragraph("Representative: Nguyen Tung Lam", style="IndentedStyle")
    doc.add_paragraph("Mobile Phone: +84909823456", style="IndentedStyle")
    doc.add_paragraph("Email: nguyentunglam@kinghospital.com", style="IndentedStyle")

    doc.add_heading(f'Party B: {contract.client.name}', level=2)
    doc.add_paragraph(f"Address: {contract.client.address}", style='IndentedStyle')
    doc.add_paragraph(f"Representative: {contract.client.representative.name}", style="IndentedStyle")
    doc.add_paragraph(f"Mobile Phone: {contract.client.representative.phone}", style="IndentedStyle")
    doc.add_paragraph(f"Email: {contract.client.representative.email}", style="IndentedStyle")
    
    # Recitals
    doc.add_heading('RECITALS')

    doc.add_paragraph("Party A have appropriate legal right, technology and resource to perform annual health check for organizations", style="IndentedStyle")
    doc.add_paragraph("Party B need to provide their employees the benefit of employee annual health check", style="IndentedStyle")

    # Agreements
    doc.add_heading('AGREEMENTS')

    doc.add_paragraph("Party A will perform health check service for party B's employees regarding the attached list with the chosen examinations and tests", style="IndentedStyle")

    doc.add_heading("PARTY B'S EMPLOYEE QUANTITY", level=2)

    doc.add_paragraph(f"Male employees: {contract.male_headcount}", style="IndentedStyle")

    doc.add_paragraph(f"Female employees: {contract.female_headcount}", style="IndentedStyle")

    # Services to be performed
    doc.add_heading('SERVICES TO BE PROVIDED', level=2)

    doc.add_paragraph('1. TEST OR EXAMINATION TO BE PERFORMED', style='IndentedStyle1')

    for service in contract.services.all():
        doc.add_paragraph({service.name}, style="IndentedStyle")    
    
    doc.add_paragraph('2. EACH PERSONNEL HEALTH REPORT WAS SENT TO EACH EMPLOYEE SEALED IN AN ENVELOPE', style='IndentedStyle1')

    doc.add_paragraph('3. GENERAL EMPLOYEES HEALTH CLASSIFYING WAS SENT TO THE HUMAN RESOURCE DEPARTMENT OF PARTY B', style='IndentedStyle1')

    doc.add_heading('CONTRACT VALUE', level=2)

    # Set the locale
    locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')

    total_value_with_separator = locale.format_string("%.2f", contract.total_value, grouping=True)
    doc.add_paragraph(f"Total Value: {total_value_with_separator} VND", style="IndentedStyle")

    doc.add_paragraph(f"Discount: {contract.discount}%", style="IndentedStyle")

    revenue_with_separator = locale.format_string("%.2f", contract.revenue, grouping=True)
    doc.add_paragraph(f"Value after discount: {revenue_with_separator} VND", style="IndentedStyle")

    doc.add_heading('CONTRACT INITIATION', level=2)

    doc.add_paragraph(f"Party A will start to perform health check for party B with the employees' detail, quantity and name attached from the initiation date to the final date in the duration of the agreement", style="IndentedStyle")

    doc.add_paragraph(f"Initiation Date: {contract.initiation_date}", style="IndentedStyle")

    doc.add_paragraph(f"Duration: 7 working days after initiation date", style="IndentedStyle")

    doc.add_heading ('PAYMENT DURATION', level=2)
    doc.add_paragraph("30 days after party A finished performing health check, send summarize health report, and payment request to party B", style="IndentedStyle")

    doc.add_heading('RESPONSIBILITY', level=2)

    doc.add_heading('PARTY A', level=3)

    doc.add_paragraph("Prepare sufficient resources including machines, devices, doctors, nurses, and other supporting staff to successfully perform health check regarding people quantity in the contract.", style="IndentedStyle")

    doc.add_paragraph("Organize the event, give the clear instruction, and inform the capability of the performance team so that party B can organize their work flow while executing the annual health check successfully, and inform their employees to plan their coming ahead", style="IndentedStyle")

    doc.add_heading('PARTY B', level=3)

    doc.add_paragraph("Inform employee about the time and place to perform health check, and to cooperate with performing team to have best efficiency", style="IndentedStyle")

    doc.add_paragraph("Pay the value due to the value in the contract, plus the additional fee in case the real quantity greater than the quantity in the contract", style="IndentedStyle")

    # Add table with one row and 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set the width of each column
    table.columns[0].width = Inches(7)
    table.columns[1].width = Inches(2)

    # Access the cells in the first row
    cells = table.rows[0].cells

    # Add content to the left cell
    left_cell = cells[0]
    left_cell.text = "Party A to sign"
    left_cell.paragraphs[0].runs[0].bold = True
    # Add indentation to left cell
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.paragraph_format.left_indent = Inches(0.5)

    # Add content to the right cell
    right_cell = cells[1]
    right_cell.text = "Party B to sign"
    right_cell.paragraphs[0].runs[0].bold = True
    
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = f'attachment; filename=contract_{contract.id}.docx'

    # save the doc
    doc.save(response)

    return response


# Allow team to manage timeline, schedule, meeting to meet up with clients
@login_required
def add_meeting(request):
    # If user submitted form
    if request.method == 'POST':
        # Check if user has permission to modify meeting
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the right to add or modify meeting"
            return HttpResponseRedirect(reverse('index'))

        # Get input from user input
        client_id = request.POST.get('client_id')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')

        if not start_time or not end_time:
            request.session['nay_message'] = 'You have to choose meeting time'
            return HttpResponseRedirect(reverse('add_meeting'))
        
        if not client_id:
            request.session['nay_message'] = 'You have to choose a client'
            return HttpResponseRedirect(reverse('add_meeting'))

        # Check if client exist
        try:
            client = Company.objects.get(pk=client_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Client does not exist"
            return HttpResponseRedirect(reverse('add_meeting'))

        # Transform start and end time from string to datetime
        start_datetime = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
        end_datetime = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')

        # Check if end time come before start time
        if start_datetime >= end_datetime:
            request.session['nay_message'] = "End time have to come after Start time"
            return HttpResponseRedirect(reverse('add_meeting'))

        newMeetUp = MeetUp(
            start_time = start_datetime,
            end_time = end_datetime,
            client = client
        )

        # Check if meeting overlap each other
        overlapping_meetings = MeetUp.objects.filter(
            start_time__lt = newMeetUp.end_time,
            end_time__gt = newMeetUp.start_time
        )

        if overlapping_meetings.exists():
            request.session['nay_message'] = 'New meeting overlapped other meeting'
            return HttpResponseRedirect(reverse('add_meeting'))

        newMeetUp.save()
        request.session['yay_message'] = 'Meeting added'

        return HttpResponseRedirect(reverse('upcoming_meetings'))

    # If user clicked link or being redirected
    else:
        # Check if user has the permission to add or modify meeting
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the permission to access this page"
            return HttpResponseRedirect(reverse('index'))

        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        companies = Company.objects.all()
        return render(request, "clinic/add_meeting.html", {
            "companies": companies,
            "yay_message": yay_message,
            "nay_message": nay_message
        })
    

# Allow user to tract all the meeting
@login_required
def all_meetings(request):
    # Check if user has permission to read meeting information
    if "read_meeting_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the permission to read this information"
        return HttpResponseRedirect(reverse('index'))

    meetings = MeetUp.objects.all().order_by("-start_time")

    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    return render(request, "clinic/all_meetings.html", {
        "meetings": meetings,
        "yay_message": yay_message,
        "nay_message": nay_message
    })


# Allow user to access list of all upcoming meeting
@login_required
def upcoming_meetings(request):
    #Check if user has the permission to read meeting information
    if "read_meeting_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the permission to read this information"
        return HttpResponseRedirect(reverse('index'))

    now = datetime.now()
    meetings = MeetUp.objects.filter(start_time__gt = now).order_by("start_time")

    # Try to warn user if meeting is close
    today = date.today()
    for meeting in meetings:
        if days_between(str(today), str(meeting.start_time.date())) <= 3:
            meeting.in_within_3_days = True
        else:
            meeting.in_within_3_days = False

    return render(request, "clinic/upcoming_meetings.html", {
        "meetings": meetings
    })


# Allow user to access meeting agenda
@login_required
def meeting_agenda(request, meeting_id):
    # If user does not have the permission to read meeting information, redirect to index page
    if "read_meeting_info" not in user_right(request.user.management_right_level):
        request.session['nay_message'] = "You do not have the permission to check this information"
        return HttpResponseRedirect(reverse('index'))

    meeting = MeetUp.objects.get(pk=meeting_id)

    # Check if user has the permission to modify meeting information or not, if not, hide the form and button in html file
    if "modify_meeting_info" not in user_right(request.user.management_right_level):
        meeting.edit_permission = False
    else:
        meeting.edit_permission = True

    yay_message = request.session.get('yay_message', '')
    nay_message = request.session.get('nay_message', '')
    request.session['yay_message'] = ''
    request.session['nay_message'] = ''

    return render(request, "clinic/meeting_agenda.html", {
        "meeting": meeting,
        "nay_message": nay_message,
        "yay_message": yay_message
    })


# allow user to add item to meeting's agenda
@login_required
def add_meeting_agenda(request, meeting_id):
    # if user submited form
    if request.method == 'POST':
        # Check user permission
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have permission to do this"
            return HttpResponseRedirect(reverse('index'))

        try:
            meeting = MeetUp.objects.get(pk=meeting_id)
        except MeetUp.DoesNotExist:
            request.session['nay_message'] = "Meeting id does not exist"
            return HttpResponseRedirect(reverse('upcoming_meetings'))
        
        # If the meetup already over, return error, not user end time because real meeting may extend longer than the preset end time
        if meeting.end_or_not:
            request.session['nay_message'] = "Meeting already over"
            return redirect('meeting_agenda', meeting_id=meeting.id)

        # Get the item
        items = {}
        for i in range(1, 6):
            if (f"item_{i}") in request.POST:
                items[f"item_{i}"] = request.POST[f"item_{i}"] or ""

                if items[f"item_{i}"] != "":

                    new_item = MeetingAgendaItem(
                        meetup = meeting,
                        item = items[f"item_{i}"]
                    )
                    new_item.save()

        # If no item was added, return nay message and redirect back
        if all(value == "" for value in items.values()):
            request.session['nay_message'] = "Please enter at least one item"
            return redirect('meeting_agenda', meeting_id=meeting_id)

        # If exist item, return yay message, then redirect also
        request.session['yay_message'] = "Item added"
        return redirect('meeting_agenda', meeting_id=meeting.id)

    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('index'))
    

# Allow user to remove item in meeting agenda
@login_required
def meeting_item_remove(request, meeting_id):
    # If user submitted form
    if request.method == 'POST':
        #Check user permission
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the permission to edit meeting information"
            return HttpResponseRedirect(reverse('index'))
        
        # Check if meeting exist
        try:
            meeting = MeetUp.objects.get(pk=meeting_id)
        except MeetUp.DoesNotExist:
            request.session['nay_message'] = "Meeting id does not exist"
            return HttpResponseRedirect(reverse('index'))
        
        # Check if meeting was ended
        if meeting.end_or_not:
            request.session['nay_message'] = "Can't remove item from ended meeting"
            return redirect("meeting_agenda", meeting_id=meeting_id)

        # Get the item's id
        item_id = request.POST.get('item_id', '')
        if not item_id:
            request.session['nay_message'] = "Item id not found"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        try: 
            item = MeetingAgendaItem.objects.get(pk=item_id)
        except MeetingAgendaItem.DoesNotExist:
            request.session['nay_message'] = "Item not found"
            return redirect("meeting_agenda", meeting_id=meeting_id)

        # Delete the item from database
        item.delete()

        # Confirm result and redirect user back
        request.session['yay_message'] = "Item removed"
        return redirect("meeting_agenda", meeting_id=meeting_id)

    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('all_meetings'))


# Allow user to modify meeting information
@login_required
def edit_meeting(request, meeting_id):
    if request.method == 'POST':
        # Check user permission
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have permission to edit meeting information"
            return HttpResponseRedirect(reverse('index'))

        try:
            meeting = MeetUp.objects.get(pk=meeting_id)
        except MeetUp.DoesNotExist:
            request.session['nay_message'] = "Meeting id not found"
            return HttpResponseRedirect(reverse('all_meetings'))
        
        # If meeting was ended, return error
        if meeting.end_or_not:
            request.session['nay_message'] = "Ended meeting can't be edited"
            return redirect("meeting_agenda", meeting_id=meeting_id)

        # Get the variables
        client_id = request.POST.get('client_id', '')
        if not client_id:
            request.session['nay_message'] = "Client id missng"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        try:
            client = Company.objects.get(pk=client_id)
        except Company.DoesNotExist:
            request.session['nay_message'] = "Client does not exist"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        
        start_time = request.POST.get('start_time', '')
        end_time = request.POST.get('end_time', '')

        if not start_time or not end_time:
            request.session['nay_message'] = "Start and end time required"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        
        try:
            start_time = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
            end_time = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
        except ValueError:
            request.session['nay_message'] = "Invalid time format"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        
        # Check if end time come before start time
        if start_time >= end_time:
            request.session['nay_message'] = "End time have to come after Start time"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        
        # Check if meeting overlap each other, exclude itself (that's mean the meeting A overlap itself is okay)
        overlapping_meetings = MeetUp.objects.filter(
            start_time__lt = end_time,
            end_time__gt = start_time
        ).exclude(id=meeting_id)

        # if exist a meeting that overlap (but not itself), return error
        if overlapping_meetings.exists():
            request.session['nay_message'] = "Modified meeting overlapped other meeting"
            return redirect("meeting_agenda", meeting_id=meeting_id)
        
        # Save meeting
        meeting.client = client
        meeting.start_time = start_time
        meeting.end_time = end_time
        meeting.save()

        # Inform successfully and redirect user
        request.session['yay_message'] = "Meeting modified"
        return redirect("meeting_agenda", meeting_id=meeting_id)
    
    else:
        request.session['nay_message'] = "POST method is required by the system"
        return HttpResponseRedirect(reverse('index'))


# Allow user to end a meeting
@login_required
def end_meeting(request):
    if request.method == 'POST':
        # Check user permission
        if "modify_meeting_info" not in user_right(request.user.management_right_level):
            request.session['nay_message'] = "You do not have the permission to end meeting"
            return HttpResponseRedirect(reverse('index'))

        meeting_id = request.POST.get('meeting_id', '')

        try:
            meeting = MeetUp.objects.get(pk=meeting_id)
        except MeetUp.DoesNotExist:
            request.session['nay_message'] = "Meeting id does not exist"

        if not meeting.end_or_not:
            # Get the meeting's agenda
            meeting_items = meeting.agenda.all()

            # Get the agenda results from POST
            for item in meeting_items:
                result = request.POST.get(f"item_result_{ item.id }")
                if not result:
                    request.session['nay_message'] = f"Result of this item was missing: {item.item}"
                    return redirect("meeting_agenda", meeting_id=meeting_id)
                item.result = result
            
            # If all the results were provided, save all of them
            for item in meeting_items:
                item.save()

            meeting.end_or_not = True
        else:
            request.session['nay_message'] = "Meeting was ended already"
            return redirect('meeting_agenda', meeting_id=meeting.id)

        meeting.save()
        request.session['yay_message'] = "Meeting set to ended"

        return redirect('meeting_agenda', meeting_id=meeting.id)
    
    else:
        request.session['nay_message'] = "POST method required"
        return HttpResponseRedirect(reverse('index'))
    

# Allow user to search in service page, people page, contract page, meeting page
@login_required
def search(request):
    # If user submit form
    if request.method == 'POST':
        # Get the variable
        search_field = request.POST.get('search-field')
        search_value = request.POST.get('search-value')
        
        # If user want to search service
        if search_field == 'service':                
            services = Service.objects.all()
            match = []
            for service in services:
                if search_value.lower() in service.name.lower():
                    match.append(service)

            return render(request, "clinic/search_result.html", {
                "match": match,
                "search_field": search_field,
                "search_value": search_value
            })

        elif search_field == 'people':
            people = People.objects.all()
            match = []
            for person in people:
                if search_value.lower() in person.name.lower():
                    match.append(person)

            return render(request, "clinic/search_result.html", {
                "match": match,
                "search_field": search_field,
                "search_value": search_value
            })

        elif search_field == 'companies':
            companies = Company.objects.all()
            match = []
            for company in companies:
                if search_value.lower() in company.name.lower():
                    match.append(company)

            return render(request, "clinic/search_result.html", {
                "match": match,
                "search_field": search_field,
                "search_value": search_value
            })

        elif search_field == 'meetings':
            meetings = MeetUp.objects.all()
            match = []
            for meeting in meetings:
                if search_value.lower() in meeting.client.name.lower():
                    match.append(meeting)

            return render(request, "clinic/search_result.html", {
                "match": match,
                "search_field": search_field,
                "search_value": search_value
            })

        elif search_field == 'contracts':
            contracts = Contract.objects.all()
            match = []
            for contract in contracts:
                if search_value.lower() in contract.client.name.lower():
                    match.append(contract)

            return render(request, "clinic/search_result.html", {
                "match": match,
                "search_field": search_field,
                "search_value": search_value
            })            

        else:
            request.session['nay_message'] = "Invalid credentials"
            return HttpResponseRedirect(reverse('search'))
    # If user clicked link or being redirected
    else:
        yay_message = request.session.get('yay_message', '')
        nay_message = request.session.get('nay_message', '')
        request.session['yay_message'] = ''
        request.session['nay_message'] = ''

        return render(request, "clinic/search.html", {
            "yay_message": yay_message,
            "nay_message": nay_message
        })

